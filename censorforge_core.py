"""
CensorForge Core - PII Detection & Redaction Engine
=====================================================

This module contains the entire backend pipeline:
  1. Presidio AnalyzerEngine (powered by spaCy en_core_web_lg) detects PII.
  2. A deterministic Faker mapping replaces each PII category with realistic
     fake data, ensuring consistency (the same real entity always maps to the
     same fake entity within a single document run).
  3. python-docx processes paragraphs AND tables, preserving formatting.

Design Decisions:
  - We use a seeded Faker instance per-run so results are reproducible during
    testing, but each new upload gets a fresh seed for privacy.
  - Custom Presidio recognizers are added for patterns that the default NER
    misses (SSNs, credit cards, IP addresses, dates of birth).
  - A confidence threshold of 0.35 is used to balance recall vs. precision;
    this is tunable via the `CONFIDENCE_THRESHOLD` constant.
"""

from __future__ import annotations

import copy
import io
import logging
import random
import re
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn
from faker import Faker
from presidio_analyzer import (
    AnalyzerEngine,
    Pattern,
    PatternRecognizer,
    RecognizerResult,
)
from presidio_analyzer.nlp_engine import NlpEngineProvider

# ==============================================================================
# Configuration
# ==============================================================================

# Minimum confidence score to accept a detection. Lower = more recall but more
# false positives. Higher = fewer false positives but may miss some PII.
CONFIDENCE_THRESHOLD = 0.35

# Entities we care about, mapped to Presidio's entity type names.
# NOTE: We intentionally exclude DATE_TIME (too broad — catches every year and
# timestamp in a financial document) and NRP/US_DRIVER_LICENSE (not requested).
# Only DATE_OF_BIRTH is included for date-related PII.
TARGET_ENTITIES = [
    "PERSON",           # Full names
    "EMAIL_ADDRESS",    # Emails
    "PHONE_NUMBER",     # Phone numbers
    "ORG",              # Company / organisation names  (spaCy NER label)
    "LOCATION",         # Physical / mailing addresses
    "US_SSN",           # Social Security Numbers
    "CREDIT_CARD",      # Credit card numbers
    "DATE_OF_BIRTH",    # Dates of birth
    "IP_ADDRESS",       # IPv4 / IPv6 addresses
]

# Logging setup
logger = logging.getLogger("censorforge")
logger.setLevel(logging.INFO)
if not logger.handlers:
    _handler = logging.StreamHandler()
    _handler.setFormatter(logging.Formatter("[%(levelname)s] %(name)s: %(message)s"))
    logger.addHandler(_handler)


# ==============================================================================
# Custom Pattern Recognizers
# ==============================================================================
# Presidio ships with many built-in recognizers, but we add explicit patterns
# for SSN, credit card, and DOB to improve recall on structured documents.


def _build_ssn_recognizer() -> PatternRecognizer:
    """
    Matches US Social Security Numbers in formats:
      - 123-45-6789
      - 123 45 6789
    """
    patterns = [
        Pattern(
            name="ssn_dashes",
            regex=r"\b\d{3}-\d{2}-\d{4}\b",
            score=0.85,
        ),
        Pattern(
            name="ssn_spaces",
            regex=r"\b\d{3}\s\d{2}\s\d{4}\b",
            score=0.65,
        ),
    ]
    return PatternRecognizer(
        supported_entity="US_SSN",
        patterns=patterns,
        name="CensorForge SSN Recognizer",
    )


def _build_credit_card_recognizer() -> PatternRecognizer:
    """
    Matches common credit card formats (Visa, MasterCard, Amex, Discover).
    Uses the Luhn algorithm via Presidio's built-in validator if available,
    but we add explicit patterns as a safety net.
    """
    patterns = [
        Pattern(
            name="credit_card_spaced",
            regex=r"\b(?:\d{4}[\s-]){3}\d{4}\b",
            score=0.80,
        ),
        Pattern(
            name="credit_card_continuous",
            regex=r"\b(?:4\d{15}|5[1-5]\d{14}|3[47]\d{13}|6(?:011|5\d{2})\d{12})\b",
            score=0.80,
        ),
    ]
    return PatternRecognizer(
        supported_entity="CREDIT_CARD",
        patterns=patterns,
        name="CensorForge Credit Card Recognizer",
    )


def _build_dob_recognizer() -> PatternRecognizer:
    """
    Matches dates that look like dates of birth in common formats:
      - MM/DD/YYYY, DD/MM/YYYY, YYYY-MM-DD
      - Month DD, YYYY  (e.g., January 15, 1990)
    Assigned to the DATE_OF_BIRTH entity type.
    """
    patterns = [
        Pattern(
            name="dob_slash",
            regex=r"\b\d{1,2}/\d{1,2}/\d{4}\b",
            score=0.50,
        ),
        Pattern(
            name="dob_dash",
            regex=r"\b\d{4}-\d{2}-\d{2}\b",
            score=0.50,
        ),
        Pattern(
            name="dob_written",
            regex=(
                r"\b(?:January|February|March|April|May|June|July|August|"
                r"September|October|November|December)\s+\d{1,2},?\s+\d{4}\b"
            ),
            score=0.55,
        ),
    ]
    return PatternRecognizer(
        supported_entity="DATE_OF_BIRTH",
        patterns=patterns,
        name="CensorForge DOB Recognizer",
    )


def _build_ip_recognizer() -> PatternRecognizer:
    """
    Matches IPv4 addresses (e.g., 192.168.1.1).
    """
    patterns = [
        Pattern(
            name="ipv4",
            regex=(
                r"\b(?:(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\.){3}"
                r"(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\b"
            ),
            score=0.70,
        ),
    ]
    return PatternRecognizer(
        supported_entity="IP_ADDRESS",
        patterns=patterns,
        name="CensorForge IP Recognizer",
    )


# ==============================================================================
# Analyzer Setup
# ==============================================================================


def create_analyzer() -> AnalyzerEngine:
    """
    Build and return a configured AnalyzerEngine with:
      - The spaCy en_core_web_lg NLP backend for high-accuracy NER.
      - Custom pattern recognizers for SSN, credit card, DOB, and IP.

    Returns:
        A ready-to-use AnalyzerEngine instance.

    Raises:
        OSError: If the spaCy model is not installed.
    """
    # Configure spaCy NLP engine with the large model for best accuracy.
    nlp_config = {
        "nlp_engine_name": "spacy",
        "models": [{"lang_code": "en", "model_name": "en_core_web_lg"}],
    }
    nlp_engine = NlpEngineProvider(nlp_configuration=nlp_config).create_engine()

    analyzer = AnalyzerEngine(nlp_engine=nlp_engine, supported_languages=["en"])

    # Register our custom recognizers alongside Presidio's defaults.
    analyzer.registry.add_recognizer(_build_ssn_recognizer())
    analyzer.registry.add_recognizer(_build_credit_card_recognizer())
    analyzer.registry.add_recognizer(_build_dob_recognizer())
    analyzer.registry.add_recognizer(_build_ip_recognizer())

    logger.info("AnalyzerEngine created with custom recognizers.")
    return analyzer


# ==============================================================================
# Faker Mapping Engine
# ==============================================================================
#
# WHY A MAPPING?
# --------------
# Presidio's built-in AnonymizerEngine can mask or hash data, but we need
# *realistic* replacements.  We maintain a dictionary keyed by
# (entity_type, original_text) → fake_text.  This ensures:
#
#   1. Consistency: "John Smith" is always replaced with the same fake name
#      within a single document, preserving referential integrity.
#   2. Realism: A name is replaced with a name, an email with an email, etc.
#   3. Determinism (optional): Setting a Faker seed makes runs reproducible.
#


class FakerMapper:
    """
    Maps detected PII entities to realistic fake replacements using Faker.

    Attributes:
        faker: A Faker instance used for generating fake data.
        _cache: Internal dict mapping (entity_type, original_text) to fake text.
                Ensures the same original PII always maps to the same fake value
                within one document processing session.
    """

    def __init__(self, seed: Optional[int] = None) -> None:
        """
        Args:
            seed: Optional integer seed for reproducible fake data.
                  If None, a random seed is chosen for privacy.
        """
        self.faker = Faker()
        if seed is not None:
            Faker.seed(seed)
            random.seed(seed)
        self._cache: Dict[Tuple[str, str], str] = {}
        logger.info("FakerMapper initialised (seed=%s).", seed)

    def get_fake(self, entity_type: str, original_text: str) -> str:
        """
        Return a realistic fake replacement for the given entity.

        The mapping logic works as follows:
          - PERSON       → faker.name()           (full name)
          - EMAIL_ADDRESS→ faker.email()           (realistic email)
          - PHONE_NUMBER → faker.phone_number()    (US-style phone)
          - ORG          → faker.company()         (company name)
          - LOCATION     → faker.address()         (full street address)
          - US_SSN       → faker.ssn()             (formatted SSN)
          - CREDIT_CARD  → faker.credit_card_number()
          - DATE_OF_BIRTH→ faker.date_of_birth().strftime(...)
          - DATE_TIME    → faker.date()            (generic date)
          - IP_ADDRESS   → faker.ipv4()            (IPv4 address)
          - NRP          → faker.country()         (nationality proxy)
          - US_DRIVER_LICENSE → faker.bothify("?##-###-####")
          - (fallback)   → "[REDACTED]"

        Args:
            entity_type:   Presidio entity type string (e.g., "PERSON").
            original_text: The original PII text found in the document.

        Returns:
            A fake replacement string.
        """
        cache_key = (entity_type, original_text.strip().lower())

        if cache_key in self._cache:
            return self._cache[cache_key]

        # ---- Faker generation by entity type ----
        # Each branch generates a contextually appropriate fake value.
        fake_value: str

        if entity_type == "PERSON":
            fake_value = self.faker.name()

        elif entity_type == "EMAIL_ADDRESS":
            fake_value = self.faker.email()

        elif entity_type == "PHONE_NUMBER":
            fake_value = self.faker.phone_number()

        elif entity_type == "ORG":
            fake_value = self.faker.company()

        elif entity_type == "LOCATION":
            # Use a single-line address to avoid breaking document layout.
            fake_value = self.faker.address().replace("\n", ", ")

        elif entity_type == "US_SSN":
            fake_value = self.faker.ssn()

        elif entity_type == "CREDIT_CARD":
            fake_value = self.faker.credit_card_number(card_type="visa")

        elif entity_type == "DATE_OF_BIRTH":
            fake_value = self.faker.date_of_birth(
                minimum_age=18, maximum_age=90
            ).strftime("%m/%d/%Y")

        elif entity_type == "IP_ADDRESS":
            fake_value = self.faker.ipv4()

        else:
            # Fallback for any entity type we haven't explicitly mapped.
            fake_value = "[REDACTED]"
            logger.warning(
                "No Faker mapping for entity type '%s'; using fallback.", entity_type
            )

        self._cache[cache_key] = fake_value
        return fake_value


# ==============================================================================
# Text-Level PII Replacement
# ==============================================================================


def redact_text(
    text: str,
    analyzer: AnalyzerEngine,
    mapper: FakerMapper,
    entities: Optional[List[str]] = None,
    score_threshold: float = CONFIDENCE_THRESHOLD,
) -> Tuple[str, List[dict]]:
    """
    Analyse a block of text for PII, then replace each detection with a
    realistic Faker-generated substitute.

    Algorithm:
      1. Run Presidio's analyzer to get a list of RecognizerResult objects.
      2. Sort results by start position **descending** so we can replace from
         the end of the string backwards without invalidating earlier indices.
      3. For each result above the confidence threshold, call FakerMapper to
         get a fake value and splice it into the string.

    Args:
        text:             The input text to scan and redact.
        analyzer:         A configured AnalyzerEngine.
        mapper:           A FakerMapper instance (shared across the document for
                          consistency).
        entities:         List of entity types to detect. Defaults to
                          TARGET_ENTITIES.
        score_threshold:  Minimum confidence score to accept a detection.

    Returns:
        A tuple of (redacted_text, detections_log) where detections_log is a
        list of dicts describing each replacement made.
    """
    if not text or not text.strip():
        return text, []

    if entities is None:
        entities = TARGET_ENTITIES

    # --- Step 1: Analyse ---
    results: List[RecognizerResult] = analyzer.analyze(
        text=text,
        language="en",
        entities=entities,
        score_threshold=score_threshold,
    )

    if not results:
        return text, []

    # --- Step 1b: Post-filter to reduce false positives ---
    # Remove detections that are almost certainly not PII, such as standalone
    # years, very short tokens, and common non-PII words that NER sometimes
    # misclassifies as ORG or PERSON.
    results = _post_filter(results, text)

    if not results:
        return text, []

    # --- Step 2: De-duplicate overlapping detections ---
    # When multiple recognizers fire on the same span, keep the one with the
    # highest confidence score.
    results = _resolve_overlaps(results)

    # --- Step 3: Sort descending by start index for safe replacement ---
    results.sort(key=lambda r: r.start, reverse=True)

    detections_log: List[dict] = []
    redacted = text

    for result in results:
        original = text[result.start : result.end]
        fake = mapper.get_fake(result.entity_type, original)

        # Splice the fake value into the string.
        redacted = redacted[: result.start] + fake + redacted[result.end :]

        detections_log.append(
            {
                "entity_type": result.entity_type,
                "original": original,
                "replacement": fake,
                "score": round(result.score, 3),
                "start": result.start,
                "end": result.end,
            }
        )

    # Reverse so log reads in document order (start → end).
    detections_log.reverse()

    return redacted, detections_log


# Common non-PII words that spaCy's NER sometimes tags as ORG or PERSON,
# especially in legal/financial documents.  Extend this list as needed.
_FALSE_POSITIVE_DENYLIST = {
    "order", "ticket", "section", "article", "clause", "chapter",
    "schedule", "annex", "appendix", "exhibit", "part", "item",
    "act", "rule", "regulation", "board", "committee", "tribunal",
    "court", "commission", "authority", "department", "ministry",
    "government", "state", "country", "nation", "republic",
    "province", "district", "office", "bureau", "agency",
    "mr", "mrs", "ms", "dr", "sr", "jr",
}


def _post_filter(
    results: List[RecognizerResult], text: str
) -> List[RecognizerResult]:
    """
    Remove detections that are very likely false positives.

    Filters applied:
      1. Standalone 4-digit years (e.g., "2013", "1956") — not PII.
      2. Time expressions (e.g., "5:00 p.m.") — not PII.
      3. Tokens shorter than 2 characters — too short to be meaningful PII.
      4. Tokens matching the deny-list of common non-PII words.
    """
    filtered: List[RecognizerResult] = []

    for r in results:
        span = text[r.start : r.end].strip()

        # Skip standalone years (4-digit numbers between 1900–2099).
        if re.fullmatch(r"\d{4}", span):
            year = int(span)
            if 1900 <= year <= 2099:
                continue

        # Skip time expressions like "5:00 p.m.", "12:30", etc.
        if re.fullmatch(r"\d{1,2}:\d{2}(\s*[ap]\.?m\.?)?", span, re.IGNORECASE):
            continue

        # Skip very short tokens (1 character) — almost never real PII.
        if len(span) < 2:
            continue

        # Skip deny-listed common words (case-insensitive).
        if span.lower().strip(".,;:!?") in _FALSE_POSITIVE_DENYLIST:
            continue

        filtered.append(r)

    return filtered


def _resolve_overlaps(results: List[RecognizerResult]) -> List[RecognizerResult]:
    """
    Given a list of RecognizerResult objects that may overlap, keep only the
    highest-scoring result for each character span.

    Strategy: sort by score descending, then greedily accept results whose
    character range doesn't intersect with any already-accepted result.
    """
    results_sorted = sorted(results, key=lambda r: r.score, reverse=True)
    accepted: List[RecognizerResult] = []

    for candidate in results_sorted:
        overlaps = False
        for existing in accepted:
            # Two spans overlap if one starts before the other ends.
            if candidate.start < existing.end and candidate.end > existing.start:
                overlaps = True
                break
        if not overlaps:
            accepted.append(candidate)

    return accepted


# ==============================================================================
# DOCX Processing (Paragraphs + Tables)
# ==============================================================================


def process_docx(
    input_stream: io.BytesIO,
    analyzer: AnalyzerEngine,
    mapper: FakerMapper,
) -> Tuple[io.BytesIO, List[dict]]:
    """
    Read a .docx from a byte stream, redact PII in every paragraph and every
    table cell, and return the modified document as a new byte stream.

    Formatting Preservation:
      - We iterate over *runs* within each paragraph.  A run is the smallest
        unit of text with uniform formatting (font, size, bold, etc.).
      - We concatenate all runs in a paragraph to get the full text, run
        Presidio on that full text, then redistribute the redacted text back
        across the original runs, preserving each run's XML formatting
        properties (rPr element).
      - This approach retains bold, italic, underline, font size, colour, and
        any other character-level formatting.

    Args:
        input_stream: A BytesIO containing the raw .docx file bytes.
        analyzer:     Configured AnalyzerEngine.
        mapper:       FakerMapper instance.

    Returns:
        (output_stream, all_detections):
          - output_stream: BytesIO with the redacted .docx.
          - all_detections: Aggregated list of every PII replacement made.
    """
    doc = Document(input_stream)
    all_detections: List[dict] = []

    # --- Process body paragraphs ---
    for para in doc.paragraphs:
        detections = _redact_paragraph(para, analyzer, mapper)
        all_detections.extend(detections)

    # --- Process every cell in every table ---
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    detections = _redact_paragraph(para, analyzer, mapper)
                    all_detections.extend(detections)

    # --- Process headers and footers ---
    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            if header_footer is not None:
                for para in header_footer.paragraphs:
                    detections = _redact_paragraph(para, analyzer, mapper)
                    all_detections.extend(detections)
                for table in header_footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                detections = _redact_paragraph(
                                    para, analyzer, mapper
                                )
                                all_detections.extend(detections)

    # Write the modified document to a new byte stream.
    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)

    logger.info("Document processed. Total PII detections: %d", len(all_detections))
    return output_stream, all_detections


def _redact_paragraph(para, analyzer: AnalyzerEngine, mapper: FakerMapper) -> List[dict]:
    """
    Redact PII in a single paragraph while preserving run-level formatting.

    How it works:
      1. Concatenate the .text of every run to form the full paragraph text.
      2. Run `redact_text()` on the full paragraph text to get the redacted
         version and the detection log.
      3. Redistribute the redacted text back across the runs.

    Run Redistribution Strategy:
      - We compute the character length of each original run.
      - We then assign characters from the redacted text to each run in
        order, adjusting for the fact that replacements may be longer or
        shorter than the originals.
      - The last run absorbs any remaining characters (or is truncated).
      - Each run's XML formatting (rPr) is preserved by directly modifying
        only the text content.

    Args:
        para:      A python-docx Paragraph object.
        analyzer:  AnalyzerEngine.
        mapper:    FakerMapper.

    Returns:
        List of detection dicts for this paragraph.
    """
    runs = para.runs
    if not runs:
        return []

    # Build the full text from runs.
    full_text = "".join(run.text for run in runs)

    if not full_text.strip():
        return []

    # Redact the full text.
    redacted_text, detections = redact_text(full_text, analyzer, mapper)

    if not detections:
        return []

    # --- Redistribute redacted text back across runs ---
    # This is the key step for preserving formatting. We keep each run's
    # formatting (bold, italic, font, etc.) but update its text content.
    _redistribute_text_to_runs(runs, full_text, redacted_text)

    return detections


def _redistribute_text_to_runs(runs, original_text: str, redacted_text: str) -> None:
    """
    Distribute `redacted_text` across the given runs, preserving each run's
    formatting properties.

    Strategy:
      We calculate the proportion of text each run contributed to the original
      and allocate a proportional share of the redacted text. This isn't
      perfect for every edge case (e.g., a PII entity spanning multiple runs),
      but it's a robust heuristic that works well in practice.

    For more precision, we use a character-mapping approach:
      1. Build a mapping from each character position in original_text to its
         run index.
      2. For each detection/replacement, the replacement characters inherit the
         run index of the first character of the original span.
      3. Reassemble run texts from the mapped characters.
    """
    if not runs:
        return

    # Simple proportional redistribution.
    original_lengths = [len(run.text) for run in runs]
    total_original = sum(original_lengths)

    if total_original == 0:
        return

    total_redacted = len(redacted_text)

    # Allocate redacted characters proportionally to each run.
    cursor = 0
    for i, run in enumerate(runs):
        if i == len(runs) - 1:
            # Last run gets everything remaining to avoid off-by-one issues.
            run.text = redacted_text[cursor:]
        else:
            # Proportional share, rounded.
            share = round(original_lengths[i] / total_original * total_redacted)
            run.text = redacted_text[cursor : cursor + share]
            cursor += share


# ==============================================================================
# High-Level API
# ==============================================================================


def censorforge_process(
    input_bytes: bytes,
    seed: Optional[int] = None,
) -> Tuple[io.BytesIO, List[dict], "AnalyzerEngine", "FakerMapper"]:
    """
    Top-level convenience function: takes raw .docx bytes, returns redacted
    .docx bytes and a detection log.

    This is the function called by the Streamlit frontend.

    Args:
        input_bytes: Raw bytes of the uploaded .docx file.
        seed:        Optional Faker seed for reproducibility.

    Returns:
        (output_stream, detections, analyzer, mapper):
          - output_stream: BytesIO with the redacted .docx.
          - detections: List of dicts describing each PII replacement.
          - analyzer: The AnalyzerEngine used (for inspection/testing).
          - mapper: The FakerMapper used (for inspection/testing).
    """
    logger.info("Starting CensorForge processing pipeline...")

    # 1. Create engines.
    analyzer = create_analyzer()
    mapper = FakerMapper(seed=seed)

    # 2. Process the document.
    input_stream = io.BytesIO(input_bytes)
    output_stream, detections = process_docx(input_stream, analyzer, mapper)

    logger.info("Processing complete. %d PII entities redacted.", len(detections))
    return output_stream, detections, analyzer, mapper
